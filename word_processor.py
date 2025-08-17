import docx
import streamlit as st
from io import BytesIO
import os

class WordProcessor:
    """Handles Word document (.docx) text extraction"""

    def __init__(self):
        pass

    def extract_text_from_docx(self, file_path_or_bytes):
        """
        Extract text from DOCX file.
        Accepts a file path or BytesIO object.
        """
        try:
            # Load the document
            doc = docx.Document(file_path_or_bytes)

            text_content = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text.strip())

            return "\n".join(text_content)

        except Exception as e:
            st.error(f"Error extracting text from DOCX file: {str(e)}")
            return ""

    def process_word_file(self, uploaded_file):
        """
        Process DOCX file and extract text.
        """
        try:
            file_extension = uploaded_file.name.lower().split('.')[-1]

            if file_extension != "docx":
                st.error(f"Unsupported file format: {file_extension}")
                return ""

            file_content = uploaded_file.read()
            return self.extract_text_from_docx(BytesIO(file_content))

        except Exception as e:
            st.error(f"Error processing Word file: {str(e)}")
            return ""
